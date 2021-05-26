# +
from docx import Document
from docx.shared import Inches


def i2w(save_path, rec, img1, img2, img3):

    document = Document()
    document.add_paragraph(rec)
    document.add_picture(img1, width=Inches(1.5), height=Inches(1))
    document.add_picture(img2, width=Inches(1.5), height=Inches(1))
    document.add_picture(img3, width=Inches(1.5), height=Inches(1))
    document.save(save_path)
    return

